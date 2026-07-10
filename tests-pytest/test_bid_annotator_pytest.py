"""投标文件偏离标注导出测试。"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from proofreader.exporters import annotate_bid_document
from proofreader.pipeline import Proofreader


def test_annotate_bid_document_creates_file(sample_docs_dir: Path, proofreader: Proofreader, tmp_path: Path) -> None:
    """标注函数应生成可打开的 docx 文件。"""
    result = proofreader.proofread(
        sample_docs_dir / "requirements.docx",
        sample_docs_dir / "bid.docx",
    )
    assert len(result.consistency_issues) > 0, "样例数据应存在偏离问题"

    output = tmp_path / "annotated_bid.docx"
    annotate_bid_document(result, output)

    assert output.exists()
    doc = Document(output)
    assert len(doc.paragraphs) > 0


def test_annotate_bid_document_adds_issue_notes(sample_docs_dir: Path, proofreader: Proofreader, tmp_path: Path) -> None:
    """标注文档中应包含偏离说明文字。"""
    result = proofreader.proofread(
        sample_docs_dir / "requirements.docx",
        sample_docs_dir / "bid.docx",
    )
    output = tmp_path / "annotated_bid.docx"
    annotate_bid_document(result, output)

    doc = Document(output)
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # 至少一个 issue 的说明文字应被追加到某段落
    assert any(issue.issue_id in full_text for issue in result.consistency_issues)


def test_annotate_bid_document_highlights_paragraphs(sample_docs_dir: Path, proofreader: Proofreader, tmp_path: Path) -> None:
    """偏离段落应被设置底纹。"""
    result = proofreader.proofread(
        sample_docs_dir / "requirements.docx",
        sample_docs_dir / "bid.docx",
    )
    output = tmp_path / "annotated_bid.docx"
    annotate_bid_document(result, output)

    doc = Document(output)
    highlighted_count = 0
    for paragraph in doc.paragraphs:
        pPr = paragraph._p.get_or_add_pPr()
        shd = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
        if shd is not None and shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"):
            highlighted_count += 1

    assert highlighted_count > 0


def test_annotate_bid_document_marks_red_spans(sample_docs_dir: Path, proofreader: Proofreader, tmp_path: Path) -> None:
    """偏离的具体文字应被标红加粗。"""
    from docx.shared import RGBColor

    result = proofreader.proofread(
        sample_docs_dir / "requirements.docx",
        sample_docs_dir / "bid.docx",
    )
    output = tmp_path / "annotated_bid.docx"
    annotate_bid_document(result, output)

    doc = Document(output)
    red_run_count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                red_run_count += 1

    assert red_run_count > 0


def test_annotate_bid_document_adds_word_comments(sample_docs_dir: Path, proofreader: Proofreader, tmp_path: Path) -> None:
    """生成的 docx 应包含 Word 批注 XML，且一个问题一个批注。"""
    import zipfile
    from lxml import etree

    result = proofreader.proofread(
        sample_docs_dir / "requirements.docx",
        sample_docs_dir / "bid.docx",
    )
    output = tmp_path / "annotated_bid.docx"
    annotate_bid_document(result, output)

    with zipfile.ZipFile(output, "r") as zf:
        names = zf.namelist()
        assert "word/comments.xml" in names
        comments_xml = zf.read("word/comments.xml").decode("utf-8")
        # 至少一个 issue 的批注内容应存在
        assert any(issue.issue_id in comments_xml for issue in result.consistency_issues)

        # 批注数量不应超过有 highlight_spans 的 issue 数量
        comments_root = etree.fromstring(comments_xml.encode("utf-8"))
        comment_count = len(comments_root.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment"))
        valid_issue_count = sum(
            1 for i in result.consistency_issues
            if i.bid_blocks and i.highlight_spans and any(s.strip() for s in i.highlight_spans)
        )
        assert 0 < comment_count <= valid_issue_count

        # 批注范围标记应出现在 document.xml 中
        document_xml = zf.read("word/document.xml").decode("utf-8")
        assert "w:commentRangeStart" in document_xml
        assert "w:commentRangeEnd" in document_xml
