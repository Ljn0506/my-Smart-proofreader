"""生成模拟的需求文件和投标文件，用于测试。"""
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


def _create_sample_image(path: Path):
    """创建一张包含关键字的测试截图。"""
    img = Image.new("RGB", (600, 200), color=(240, 240, 240))
    draw = ImageDraw.Draw(img)

    # 尝试使用 macOS 常见中文字体
    font_paths = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
    ]
    font = None
    for fp in font_paths:
        try:
            font = ImageFont.truetype(fp, 28)
            break
        except Exception:
            continue
    if font is None:
        font = ImageFont.load_default()

    draw.text((30, 30), "系统管理后台截图", fill=(0, 0, 0), font=font)
    draw.text((30, 80), "功能模块：数据备份、用户管理、权限控制", fill=(50, 50, 50), font=font)
    draw.text((30, 130), "系统支持每日自动备份", fill=(50, 50, 50), font=font)
    img.save(path)


def create_requirement_doc(path: Path):
    doc = Document()
    doc.add_heading("某信息系统采购项目技术需求书", level=1)

    doc.add_heading("一、项目概况", level=2)
    doc.add_paragraph("本项目采购一套信息管理系统，用于实现业务流程数字化管理。")

    doc.add_heading("二、技术要求", level=2)
    doc.add_paragraph("1. 系统必须支持 1000 并发用户同时在线访问。")
    doc.add_paragraph("2. 核心功能响应时间不得超过 2 秒。")
    doc.add_paragraph("3. 系统必须提供数据备份与恢复功能，备份周期不少于 7 天。")
    doc.add_paragraph("4. 质保期不少于 3 年，自项目终验合格之日起计算。")
    doc.add_paragraph("5. 项目经理必须具备 PMP 证书，并具有 5 年以上项目管理经验。")
    doc.add_paragraph("6. 系统必须支持国产化操作系统，包括麒麟、统信 UOS。")
    doc.add_paragraph("7. 投标方必须提供 7×24 小时技术支持服务。")

    doc.add_heading("三、交付要求", level=2)
    doc.add_paragraph("8. 项目交付周期不超过 6 个月。")
    doc.add_paragraph("9. 培训次数不少于 3 次，每次不少于 20 人。")

    doc.save(path)
    print(f"已生成需求文件：{path}")


def create_bid_doc(path: Path):
    doc = Document()
    doc.add_heading("某信息系统采购项目投标文件", level=1)

    # 商务部分
    doc.add_heading("一、商务部分", level=2)
    doc.add_paragraph("我司具有完善的项目管理体系，项目经理持有 PMP 证书。")
    doc.add_paragraph("本次投标报价为人民币 580 万元整。")

    # 技术部分
    doc.add_heading("二、技术部分", level=2)
    doc.add_paragraph("1. 系统架构设计")
    doc.add_paragraph("本系统采用微服务架购，支持高并发访问。经测试，系统可支持 800 并发用户同时在线。")
    doc.add_paragraph("核心功能响应时间为 1.5 秒，满足性能要求。")
    doc.add_paragraph("系统提供每日自动备份功能，备份数据保留 5 天。")
    doc.add_paragraph("本系统全面支持国产化操作系统，包括麒麟、统信 UOS。")
    doc.add_paragraph("技术支持服务时间为工作日 9:00-18:00。")

    doc.add_paragraph("2. 人员资质")
    doc.add_paragraph("项目经理张三具备 PMP 证书，具有 4 年项目管理经验。")

    # 插入截图
    img_path = path.parent / "sample_screenshot.png"
    _create_sample_image(img_path)
    doc.add_paragraph("3. 系统截图")
    doc.add_picture(str(img_path), width=Inches(5.0))

    # 价格部分
    doc.add_heading("三、价格部分", level=2)
    doc.add_paragraph("总报价：580 万元。")
    doc.add_paragraph("付款方式：合同签订后支付 30%，验收合格后支付 70%。")

    # 交付部分
    doc.add_heading("四、项目实施与交付", level=2)
    doc.add_paragraph("项目交付周期为 8 个月。")
    doc.add_paragraph("培训计划：提供 2 次集中培训，每次覆盖 15 人。")
    doc.add_paragraph("质保期为 2 年。")

    doc.save(path)
    print(f"已生成投标文件：{path}")


if __name__ == "__main__":
    out_dir = Path(__file__).parent.parent / "data" / "sample-docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    create_requirement_doc(out_dir / "requirements.docx")
    create_bid_doc(out_dir / "bid.docx")
