"""解析 Word .docx / .doc 文件，提取段落、标题、表格和图片位置。"""
from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


# WordprocessingML / DrawingML 命名空间
_NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


@dataclass
class TextBlock:
    """文档中的一个文本块（段落或表格行转换后的文本）。"""
    text: str
    block_type: str  # "paragraph", "table_row", "heading"
    level: int = 0  # 标题层级，0 表示非标题
    style_name: str = ""
    page_hint: int = 0  # 页码提示（docx 本身无精确页码，这里按近似估算）
    index: int = 0  # 在文档中的顺序


@dataclass
class EmbeddedImage:
    """嵌入文档的图片。"""
    image_index: int
    ext: str
    blob: bytes
    block_index: int  # 图片所在文本块序号（对应 TextBlock.index）


@dataclass
class ParsedDocument:
    """解析后的文档对象。"""
    path: Path
    blocks: List[TextBlock] = field(default_factory=list)
    images: List[EmbeddedImage] = field(default_factory=list)
    headings: List[TextBlock] = field(default_factory=list)
    raw_tables: List[List[List[str]]] = field(default_factory=list)


def _is_heading(paragraph: Paragraph) -> Tuple[bool, int]:
    """判断段落是否为标题，并返回层级。"""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading "):
        try:
            return True, int(style_name.split()[-1])
        except ValueError:
            return True, 1
    # 有些中文模板用「标题 1」「标题 2」
    if style_name.startswith("标题 "):
        try:
            return True, int(style_name.split()[-1])
        except ValueError:
            return True, 1
    return False, 0


def _cell_text(cell) -> str:
    """提取单元格文本。"""
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _table_to_text_rows(table: Table) -> List[str]:
    """把表格按行转换为文本。"""
    rows = []
    for row in table.rows:
        cells = [_cell_text(cell) for cell in row.cells]
        # 过滤空行
        if any(cells):
            rows.append(" | ".join(cells))
    return rows


def _extract_raw_table(table: Table) -> List[List[str]]:
    """保留表格原始结构。"""
    return [[_cell_text(cell) for cell in row.cells] for row in table.rows]


# LibreOffice / OpenOffice 可执行文件名（跨平台常见）
_SOFFICE_CANDIDATES = ["soffice", "libreoffice"]


def find_soffice() -> str | None:
    """在 PATH 中查找 soffice / libreoffice 命令。"""
    for candidate in _SOFFICE_CANDIDATES:
        executable = shutil.which(candidate)
        if executable:
            return executable
    return None


# 保留旧别名，兼容现有内部调用
_find_soffice = find_soffice


def convert_doc_to_docx(doc_path: Path, output_dir: Path) -> Path:
    """使用 LibreOffice 将 .doc 转换为 .docx，返回转换后的文件路径。"""
    soffice = find_soffice()
    if soffice is None:
        raise RuntimeError(
            "检测到 .doc 文件，但未找到 LibreOffice 命令（soffice / libreoffice）。"
            "请安装 LibreOffice 后重试，或先将文件另存为 .docx。"
        )

    output_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(output_dir),
        str(doc_path),
    ]
    try:
        result = subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120,
        )
    except subprocess.CalledProcessError as exc:
        stderr = exc.stderr.decode("utf-8", errors="ignore") if exc.stderr else ""
        raise RuntimeError(f"转换 .doc 文件失败：{stderr or exc}") from exc
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("转换 .doc 文件超时，请检查 LibreOffice 是否可用。") from exc

    converted = output_dir / doc_path.with_suffix(".docx").name
    if not converted.exists():
        # LibreOffice 有时会重命名输出文件
        candidates = list(output_dir.glob("*.docx"))
        if not candidates:
            raise RuntimeError("LibreOffice 转换 .doc 后未生成 .docx 文件。")
        converted = candidates[0]
    return converted


# 保留旧别名，兼容现有内部调用
_convert_doc_to_docx = convert_doc_to_docx


def _parse_docx_document(doc_path: Path, original_path: Path) -> ParsedDocument:
    """解析已转换/本身就是 .docx 的文档，返回 ParsedDocument（保留原始路径）。"""
    doc = Document(str(doc_path))

    parsed = ParsedDocument(path=original_path)
    block_index = 0
    para_index = 0
    # 记录每个顶层段落元素对应的文本块序号，用于后续图片定位
    para_element_to_block_index: Dict[object, int] = {}

    for element in doc.element.body:
        if element.tag.endswith("p"):
            paragraph = Paragraph(element, doc)
            is_heading, level = _is_heading(paragraph)
            text = paragraph.text.strip()
            if not text:
                # 空段落也可能包含图片，归到前一个文本块
                para_element_to_block_index[element] = max(0, block_index - 1)
                para_index += 1
                continue
            # 先记录段落所在块序号
            para_element_to_block_index[element] = block_index

            block = TextBlock(
                text=text,
                block_type="heading" if is_heading else "paragraph",
                level=level,
                style_name=paragraph.style.name if paragraph.style else "",
                index=block_index,
            )
            parsed.blocks.append(block)
            if is_heading:
                parsed.headings.append(block)
            block_index += 1
            para_index += 1

        elif element.tag.endswith("tbl"):
            table = Table(element, doc)
            raw_table = _extract_raw_table(table)
            if raw_table:
                parsed.raw_tables.append(raw_table)

            for row_text in _table_to_text_rows(table):
                block = TextBlock(
                    text=row_text,
                    block_type="table_row",
                    level=0,
                    style_name="Table",
                    index=block_index,
                )
                parsed.blocks.append(block)
                block_index += 1

    # 提取图片，并记录其所在的段落/块位置
    image_index = 0
    for shape in doc.inline_shapes:
        try:
            # 获取图片二进制数据
            blips = shape._inline.findall(".//a:blip", namespaces=_NSMAP)
            if not blips:
                continue
            blip = blips[0]
            embed_id = blip.get("{%s}embed" % _NSMAP["r"])
            if not embed_id:
                continue
            rel = doc.part.rels.get(embed_id)
            if rel is None or "image" not in rel.target_ref:
                continue
            image_part = rel.target_part
            ext = image_part.content_type.split("/")[-1]
            blob = image_part.blob

            # 定位图片所在段落，映射到文本块序号
            para_elements = shape._inline.xpath("ancestor::w:p")
            block_index_for_image = 0
            if para_elements:
                block_index_for_image = para_element_to_block_index.get(para_elements[0], 0)

            parsed.images.append(
                EmbeddedImage(
                    image_index=image_index,
                    ext=ext,
                    blob=blob,
                    block_index=block_index_for_image,
                )
            )
            image_index += 1
        except Exception:
            continue

    return parsed


def parse_docx(path: Path | str) -> ParsedDocument:
    """解析 Word 文件，支持 .docx 与 .doc（依赖 LibreOffice 转换）。"""
    path = Path(path)
    original_path = path

    if path.suffix.lower() == ".doc":
        with tempfile.TemporaryDirectory(prefix="smart_proofreader_doc_convert_") as tmp_dir_str:
            tmp_dir = Path(tmp_dir_str)
            converted_path = convert_doc_to_docx(path, tmp_dir)
            return _parse_docx_document(converted_path, original_path)

    return _parse_docx_document(path, original_path)
