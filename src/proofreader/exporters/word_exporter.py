"""将校对结果导出为 Word 报告。"""
from __future__ import annotations

from pathlib import Path
from typing import Any, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from proofreader.checkers.consistency_checker import ConsistencyIssue, IssueLevel
from proofreader.checkers.ocr_checker import OcrIssue
from proofreader.checkers.table_checker import TableIssue
from proofreader.checkers.typo_checker import TypoIssue
from proofreader.pipeline import ProofreadResult


def _add_heading(doc: Any, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_colored_text(paragraph, text: str, color: RGBColor, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold


def _issue_level_label(level: IssueLevel) -> str:
    return {
        IssueLevel.ERROR: "严重",
        IssueLevel.WARNING: "警告",
        IssueLevel.INFO: "提示",
    }.get(level, "未知")


def _issue_level_color(level: IssueLevel) -> RGBColor:
    return {
        IssueLevel.ERROR: RGBColor(220, 53, 69),
        IssueLevel.WARNING: RGBColor(255, 193, 7),
        IssueLevel.INFO: RGBColor(13, 202, 240),
    }.get(level, RGBColor(108, 117, 125))


def _add_consistency_issues(doc: Any, issues: List[ConsistencyIssue]) -> None:
    if not issues:
        doc.add_paragraph("未发现一致性/偏离问题。")
        return

    for issue in issues:
        p = doc.add_paragraph()
        _add_colored_text(p, f"[{_issue_level_label(issue.level)}] ", _issue_level_color(issue.level), bold=True)
        p.add_run(f"{issue.message}")

        detail = doc.add_paragraph(style="List Bullet")
        detail.add_run(f"问题类型：{issue.issue_type.value}\n").bold = True
        detail.add_run(f"需求条目：{issue.requirement_text}\n")
        detail.add_run(f"投标响应：{issue.bid_text or '无'}\n")
        detail.add_run(f"建议：{issue.suggestion}")


def _add_typo_issues(doc: Any, issues: List[TypoIssue]) -> None:
    if not issues:
        doc.add_paragraph("未发现错别字问题。")
        return

    for typo in issues:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"「{typo.word}」 → 「{typo.suggestion}」").bold = True
        p.add_run(f"（{typo.message}）")


def _add_ocr_issues(doc: Any, issues: List[OcrIssue]) -> None:
    if not issues:
        doc.add_paragraph("未发现截图 OCR 问题。")
        return

    for ocr in issues:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"图片 #{ocr.image_index}").bold = True
        if ocr.context_block:
            p.add_run(f"（位于：{ocr.context_block.text[:60]}...）")
        p.add_run(f"：{ocr.message}")


def _add_table_issues(doc: Any, issues: List[TableIssue]) -> None:
    if not issues:
        doc.add_paragraph("未发现表格比对问题。")
        return

    for table in issues:
        p = doc.add_paragraph()
        p.add_run(f"[{table.issue_id}] ").bold = True
        p.add_run(table.message)

        detail = doc.add_paragraph(style="List Bullet")
        detail.add_run(f"建议：{table.suggestion}")
        for d in table.details[:10]:
            sub = doc.add_paragraph(style="List Bullet 2")
            sub.add_run(d)


def export_to_word(result: ProofreadResult, output_path: Path | str) -> Path:
    """将校对结果导出为 Word 报告。"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 标题
    title = doc.add_heading("智能文档校对报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 摘要
    _add_heading(doc, "一、校对摘要", level=1)
    summary = doc.add_paragraph()
    summary.add_run(
        f"需求文件：{result.requirement_doc.path.name}\n"
        f"投标文件：{result.bid_doc.path.name}\n"
        f"提取需求条目：{len(result.requirements)} 条\n"
        f"一致性/偏离问题：{len(result.consistency_issues)} 个\n"
        f"表格比对问题：{len(result.table_issues)} 个\n"
        f"错别字问题：{len(result.typo_issues)} 个\n"
        f"截图 OCR 问题：{len(result.ocr_issues)} 个"
    )

    # 一致性/偏离问题
    _add_heading(doc, "二、一致性 / 偏离问题", level=1)
    _add_consistency_issues(doc, result.consistency_issues)

    # 表格比对问题
    _add_heading(doc, "三、表格比对问题", level=1)
    _add_table_issues(doc, result.table_issues)

    # 错别字
    _add_heading(doc, "四、错别字问题", level=1)
    _add_typo_issues(doc, result.typo_issues)

    # 截图 OCR
    _add_heading(doc, "五、截图 OCR 问题", level=1)
    _add_ocr_issues(doc, result.ocr_issues)

    doc.save(str(output_path))
    return output_path
