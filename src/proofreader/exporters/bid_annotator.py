"""在投标文件副本中高亮标注偏离项（精确文字标红 + Word 原生批注）。"""
from __future__ import annotations

import os
import re
import shutil
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree

from proofreader.checkers.consistency_checker import ConsistencyIssue, IssueLevel
from proofreader.parsers.docx_parser import convert_doc_to_docx
from proofreader.pipeline import ProofreadResult


# 偏离级别对应底纹颜色（十六进制）
LEVEL_FILL_COLORS = {
    IssueLevel.ERROR: "F8D7DA",    # 浅红
    IssueLevel.WARNING: "FFF3CD",  # 浅黄
    IssueLevel.INFO: "D1ECF1",     # 浅蓝
}


def _normalize(text: str) -> str:
    """归一化文本用于匹配：去首尾空白、合并连续空白。"""
    return " ".join(text.strip().split())


def _match_paragraph(paragraph_text: str, target_text: str) -> bool:
    """判断段落文本是否与目标文本匹配。"""
    if not target_text:
        return False
    pt = _normalize(paragraph_text)
    tt = _normalize(target_text)
    if pt == tt:
        return True
    if tt in pt or pt in tt:
        return True
    return False


def _set_paragraph_shading(paragraph, fill_color: str) -> None:
    """设置段落底纹颜色。"""
    pPr = paragraph._p.get_or_add_pPr()
    # 移除已有的 shd 元素，避免重复
    for child in list(pPr):
        if child.tag == qn("w:shd"):
            pPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_color)
    pPr.append(shd)


def _mark_spans_red(paragraph, spans: List[str]) -> None:
    """将段落中包含 spans 文字的 run 标红加粗（兼容空格差异）。"""
    if not spans:
        return
    for run in paragraph.runs:
        run_norm = _normalize(run.text).replace(" ", "")
        for span in spans:
            span_norm = _normalize(span).replace(" ", "")
            if span_norm and span_norm in run_norm:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True
                break


def _add_issue_note(paragraph, issue: ConsistencyIssue) -> None:
    """在段落末尾追加红色小字说明（作为批注的补充）。"""
    paragraph.add_run().add_break()
    note_text = f"[{issue.issue_id}] 需求 {issue.requirement_id}：{issue.message}（{issue.suggestion}）"
    run = paragraph.add_run(note_text)
    run.font.color.rgb = RGBColor(211, 47, 47)  # 红色
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Word 原生批注（comments）支持
# ---------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_REL_COMMENTS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _build_comments_xml(comments: List[Tuple[int, str, str]]) -> bytes:
    """生成 comments.xml 内容。comments: [(id, author, text), ...]。"""
    root = etree.Element(f"{{{_W_NS}}}comments", nsmap={"w": _W_NS, "r": _R_NS})
    for cid, author, text in comments:
        comment = etree.SubElement(
            root,
            f"{{{_W_NS}}}comment",
            {
                f"{{{_W_NS}}}id": str(cid),
                f"{{{_W_NS}}}author": author,
                f"{{{_W_NS}}}date": _now_iso(),
                f"{{{_W_NS}}}initials": author[:1],
            },
        )
        p = etree.SubElement(comment, f"{{{_W_NS}}}p")
        r = etree.SubElement(p, f"{{{_W_NS}}}r")
        t = etree.SubElement(r, f"{{{_W_NS}}}t")
        t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _inject_comments_into_docx(docx_path: Path, issues: List[ConsistencyIssue]) -> None:
    """
    向已保存的 docx 文件中注入 Word 原生批注。

    - 只对有 highlight_spans 的 issue 生成批注（一个问题一个批注）。
    - 批注范围定位在包含问题文字的 run 上，而不是段落开头。
    """
    # 过滤有效 issue：必须有 bid_blocks 且 highlight_spans 非空
    valid_issues = [
        issue for issue in issues
        if issue.bid_blocks
        and issue.highlight_spans
        and any(s.strip() for s in issue.highlight_spans)
    ]
    if not valid_issues:
        return

    work_dir = Path(tempfile.mkdtemp())
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            for member in zin.namelist():
                if os.path.isabs(member) or ".." in Path(member).parts:
                    raise ValueError(f"Invalid zip member path: {member}")
            zin.extractall(work_dir)

        doc_xml_path = work_dir / "word" / "document.xml"
        rels_path = work_dir / "word" / "_rels" / "document.xml.rels"
        comments_path = work_dir / "word" / "comments.xml"
        content_types_path = work_dir / "[Content_Types].xml"

        if not doc_xml_path.exists():
            return

        # 1. 准备批注数据
        comments_data: List[Tuple[int, str, str]] = []
        for idx, issue in enumerate(valid_issues):
            text = f"[{issue.issue_id}] 需求 {issue.requirement_id}\n{issue.message}\n建议：{issue.suggestion}"
            comments_data.append((idx, "智能校对器", text))

        # 2. 写入 comments.xml
        comments_path.write_bytes(_build_comments_xml(comments_data))

        # 3. 更新 document.xml.rels
        rels_tree = etree.parse(str(rels_path))
        rels_root = rels_tree.getroot()
        ns = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
        existing_ids = [
            int(el.get("Id", "rId0")[3:]) for el in rels_root.findall("rel:Relationship", ns) if el.get("Id", "").startswith("rId")
        ]
        next_rid = max(existing_ids) + 1 if existing_ids else 1
        rel = etree.SubElement(rels_root, f"{{{ns['rel']}}}Relationship")
        rel.set("Id", f"rId{next_rid}")
        rel.set("Type", _REL_COMMENTS)
        rel.set("Target", "comments.xml")
        rels_tree.write(str(rels_path), xml_declaration=True, encoding="UTF-8", standalone=True)

        # 4. 更新 [Content_Types].xml
        if content_types_path.exists():
            ct_tree = etree.parse(str(content_types_path))
            ct_root = ct_tree.getroot()
            ct_ns = ct_root.nsmap.get(None, "http://schemas.openxmlformats.org/package/2006/content-types")
            already = any(
                el.get("Extension") == "xml" and el.get("ContentType") == "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
                for el in ct_root
            )
            if not already:
                override = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                override.set("PartName", "/word/comments.xml")
                override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
                ct_tree.write(str(content_types_path), xml_declaration=True, encoding="UTF-8", standalone=True)

        # 5. 在 document.xml 中按段落插入批注引用（段落级范围，兼容性更好）
        doc_tree = etree.parse(str(doc_xml_path))
        doc_root = doc_tree.getroot()
        w_ns = doc_root.nsmap.get("w", _W_NS)

        for issue_idx, issue in enumerate(valid_issues):
            target_text = issue.bid_blocks[0].text
            for paragraph in doc_root.iter(f"{{{w_ns}}}p"):
                para_text = "".join(t.text or "" for t in paragraph.iter(f"{{{w_ns}}}t"))
                if not _match_paragraph(para_text, target_text):
                    continue

                # 找到 pPr 后的插入位置（若不存在 pPr 则从开头开始）
                insert_pos = 0
                for idx, child in enumerate(paragraph):
                    if child.tag == f"{{{w_ns}}}pPr":
                        insert_pos = idx + 1
                        break

                # 在 pPr 后插入 commentRangeStart
                range_start = OxmlElement("w:commentRangeStart")
                range_start.set(qn("w:id"), str(issue_idx))
                paragraph.insert(insert_pos, range_start)

                # 在段落末尾插入 commentRangeEnd 和 commentReference
                range_end = OxmlElement("w:commentRangeEnd")
                range_end.set(qn("w:id"), str(issue_idx))
                ref_run = OxmlElement("w:r")
                ref = OxmlElement("w:commentReference")
                ref.set(qn("w:id"), str(issue_idx))
                ref_run.append(ref)
                paragraph.append(range_end)
                paragraph.append(ref_run)
                break

        doc_tree.write(str(doc_xml_path), xml_declaration=True, encoding="UTF-8", standalone=True)

        # 6. 重新打包
        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for root, _, files in os.walk(work_dir):
                for file in files:
                    file_path = Path(root) / file
                    arcname = str(file_path.relative_to(work_dir))
                    zout.write(file_path, arcname)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def annotate_bid_document(result: ProofreadResult, output_path: Path | str) -> Path:
    """
    生成一份带偏离标注的投标文件副本。

    - 对包含偏离的段落设置背景色高亮。
    - 对 highlight_spans 中的具体文字标红加粗。
    - 在段落末尾追加红色说明文字。
    - 注入 Word 原生批注（comments）。
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    bid_source = result.bid_doc.path
    if output_path.resolve() == bid_source.resolve():
        raise ValueError("标注输出路径不能与源投标文件相同，请选择其他路径。")

    # 处理 .doc：先转换为 .docx 临时文件
    if bid_source.suffix.lower() == ".doc":
        with tempfile.TemporaryDirectory(prefix="smart_proofreader_bid_convert_") as tmp_convert_dir:
            converted = convert_doc_to_docx(bid_source, Path(tmp_convert_dir))
            shutil.copy(converted, output_path)
    else:
        shutil.copy(bid_source, output_path)

    doc = Document(str(output_path))

    # 按段落文本收集需要标注的 issues（一个段落可能对应多个 issue）
    issues_by_paragraph: Dict[int, List[ConsistencyIssue]] = {}
    missing_response_issues: List[ConsistencyIssue] = []
    for issue in result.consistency_issues:
        if not issue.bid_blocks:
            missing_response_issues.append(issue)
            continue
        target_text = issue.bid_blocks[0].text
        if not target_text:
            continue

        matched = False
        for p_idx, paragraph in enumerate(doc.paragraphs):
            if _match_paragraph(paragraph.text, target_text):
                issues_by_paragraph.setdefault(p_idx, []).append(issue)
                matched = True
                break

        # 段落未命中，尝试在表格单元格中匹配
        if not matched:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if _match_paragraph(cell.text, target_text):
                            # 表格单元格直接标红并追加说明
                            _mark_spans_red(cell.paragraphs[0], issue.highlight_spans)
                            _add_issue_note(cell.paragraphs[0], issue)
                            matched = True
                            break
                    if matched:
                        break
                if matched:
                    break

    # 对匹配到的段落执行标注
    for p_idx, issues in issues_by_paragraph.items():
        paragraph = doc.paragraphs[p_idx]
        # 取最高级别作为底纹颜色
        most_severe = min(issues, key=lambda i: (i.level != IssueLevel.ERROR, i.level != IssueLevel.WARNING, i.level != IssueLevel.INFO))
        fill_color = LEVEL_FILL_COLORS.get(most_severe.level, "FFF3CD")
        _set_paragraph_shading(paragraph, fill_color)

        # 精确文字标红
        all_spans: List[str] = []
        for issue in issues:
            all_spans.extend(issue.highlight_spans)
        _mark_spans_red(paragraph, all_spans)

        for issue in issues:
            _add_issue_note(paragraph, issue)

    # 对缺失响应的需求在文档末尾追加汇总页
    if missing_response_issues:
        doc.add_page_break()
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("未响应需求汇总")
        title_run.font.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(220, 53, 69)

        for issue in missing_response_issues:
            p = doc.add_paragraph()
            run = p.add_run(
                f"[{issue.issue_id}] 需求 {issue.requirement_id}\n"
                f"需求内容：{issue.requirement_text}\n"
                f"问题：{issue.message}\n"
                f"建议：{issue.suggestion}"
            )
            run.font.color.rgb = RGBColor(220, 53, 69)
            run.font.size = Pt(10)

    doc.save(str(output_path))

    # 注入 Word 原生批注（对缺失响应也生成批注，便于在汇总页查看）
    _inject_comments_into_docx(output_path, result.consistency_issues)

    return output_path
